from __future__ import annotations

import json
from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.shared import Cm, Inches, Pt


REPORT_DOCX = Path("Горкунов Н.М. ИУ5-21М ММО ДЗ loss functions.docx")
ARTIFACTS_DIR = Path("artifacts")
DATA_DIR = Path("data") / "processed"
TITLE_TEMPLATE_DOCX = Path("РНС ЛР6 ИУ5-21М Горкунов Н.М. редакция.docx")


def build_report() -> None:
    results = json.loads((ARTIFACTS_DIR / "results.json").read_text(encoding="utf-8"))
    metrics = pd.read_csv(ARTIFACTS_DIR / "metrics_summary.csv")
    dataset_summary = pd.read_csv(DATA_DIR / "summary.csv")

    document = build_document_with_title_template()
    setup_document(document)
    document.add_section(WD_SECTION.NEW_PAGE)
    add_intro(document)
    add_data_section(document, dataset_summary)
    add_method_section(document, results)
    add_results_section(document, results, metrics)
    add_conclusion(document, metrics)
    add_sources(document)
    document.save(REPORT_DOCX)


def setup_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1)

    styles = document.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(14)
    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        styles[style_name].font.name = "Times New Roman"
        styles[style_name].font.size = Pt(14)
        styles[style_name].font.bold = True
        styles[style_name].font.color.rgb = RGBColor(0, 0, 0)


def build_document_with_title_template() -> Document:
    if not TITLE_TEMPLATE_DOCX.exists():
        document = Document()
        add_title_page(document)
        return document

    document = Document(TITLE_TEMPLATE_DOCX)
    keep_title_page_only(document)
    replace_title_text(document)
    return document


def keep_title_page_only(document: Document) -> None:
    for paragraph_item in list(document.paragraphs)[24:]:
        element = paragraph_item._element
        element.getparent().remove(element)
    for table in list(document.tables)[1:]:
        element = table._element
        element.getparent().remove(element)


def replace_title_text(document: Document) -> None:
    replacements = {
        8: "ПО ДОМАШНЕМУ ЗАДАНИЮ",
        9: "По дисциплине «Методы машинного обучения»",
        11: "«Сравнение функций потерь обучения моделей эмбеддинга текстов»",
        17: "Преподаватель \t\t\t_________________\t__Ю.Е. Гапанюк__",
    }
    for paragraph_index, text in replacements.items():
        replace_paragraph_text(document.paragraphs[paragraph_index], text)


def replace_paragraph_text(paragraph_item: Any, text: str) -> None:
    if not paragraph_item.runs:
        paragraph_item.add_run(text)
        return
    paragraph_item.runs[0].text = text
    for run in paragraph_item.runs[1:]:
        run.text = ""


def add_title_page(document: Document) -> None:
    centered(document, "ФАКУЛЬТЕТ \tИНФОРМАТИКА И СИСТЕМЫ УПРАВЛЕНИЯ")
    centered(document, "КАФЕДРА \tСИСТЕМЫ ОБРАБОТКИ ИНФОРМАЦИИ И УПРАВЛЕНИЯ")
    add_spacer(document, 3)
    centered(document, "ОТЧЕТ", bold=True)
    centered(document, "ПО ДОМАШНЕМУ ЗАДАНИЮ", bold=True)
    centered(document, "По дисциплине «Методы машинного обучения»")
    add_spacer(document, 1)
    centered(document, "«Сравнение функций потерь обучения моделей эмбеддинга текстов»")
    add_spacer(document, 4)
    paragraph(document, "Студент __ИУ5-21М__\t\t_________________\t__Н.М. Горкунов__")
    small_centered(document, "(Группа)\t(Подпись, дата)\t(И.О.Фамилия)")
    paragraph(document, "Преподаватель \t\t\t_________________\t__Ю.Е. Гапанюк__")
    small_centered(document, "(Подпись, дата)\t(И.О.Фамилия)")
    add_spacer(document, 5)
    centered(document, "Москва - 2026")


def add_intro(document: Document) -> None:
    heading(document, "Введение")
    paragraph(
        document,
        "В работе сравниваются функции потерь, применяемые для обучения моделей текстовых эмбеддингов. "
        "Цель эксперимента - проверить, как разные функции потерь изменяют качество представлений при одинаковых данных, "
        "одинаковом числе обучающих записей и одинаковом количестве обучаемых параметров.",
    )
    paragraph(
        document,
        "В качестве базовой модели используется sentence-transformers/msmarco-distilbert-base-dot-prod-v3. "
        "Transformer encoder и pooling остаются замороженными, а Dense projection head размерности 768x768 "
        "переинициализируется перед каждым запуском. Такой режим изолирует влияние функции потерь лучше, чем полное "
        "дообучение всех весов модели на малом наборе данных.",
    )
    heading(document, "Задание")
    paragraph(
        document,
        "Необходимо подготовить несколько качественных датасетов для разных форматов обучения эмбеддингов, привести их "
        "к одинаковому размеру и сравнить Contrastive Loss, Triplet Loss, InfoNCE Loss, NT-Xent Loss, Supervised "
        "Contrastive Loss и Circle Loss. Результаты эксперимента фиксируются в Jupyter Notebook и используются "
        "для анализа качества функций потерь.",
    )


def add_data_section(document: Document, dataset_summary: pd.DataFrame) -> None:
    heading(document, "Подготовка данных")
    paragraph(
        document,
        "Для основного сравнения выбраны четыре англоязычных датасета: all-nli, quora-duplicates, banking77 и "
        "twenty-newsgroups. Для каждого применимого представления используется 5000 train-записей и 1000 test-записей. "
        "Датасет STS-B используется только как внешняя проверка семантической близости.",
    )
    compact = (
        dataset_summary[dataset_summary["split"].isin(["train", "test"])]
        .groupby(["dataset_id", "view", "split"], as_index=False)["rows"]
        .first()
    )
    add_table(
        document,
        "Таблица 1 - Нормализованные представления датасетов",
        compact,
        ["dataset_id", "view", "split", "rows"],
        max_rows=28,
    )
    paragraph(
        document,
        "Для pair/triplet датасетов class_text представление не строится, так как у них нет устойчивых single-text "
        "классов. Поэтому SupCon и Circle Loss сравниваются на banking77 и twenty-newsgroups.",
    )


def add_method_section(document: Document, results: dict[str, Any]) -> None:
    heading(document, "Методика эксперимента")
    config = results["config"]
    paragraph(
        document,
        f"Во всех запусках используется seed = {config['seed']}, batch_size = {config['batch_size']}, "
        f"epochs = {config['epochs']}, learning_rate = {config['learning_rate']}, weight_decay = {config['weight_decay']}. "
        f"Количество обучаемых параметров projection head равно {config['trainable_parameters']}.",
    )
    paragraph(
        document,
        "Перед обучением для всех текстов кэшируются эмбеддинги замороженной части модели: DistilBERT encoder и mean "
        "pooling. Затем для каждой функции потерь создается новая Linear projection head с Xavier-инициализацией. "
        "Именно эта матрица обучается на соответствующем представлении датасета.",
    )
    paragraph(
        document,
        f"\u0414\u043b\u044f \u0440\u0430\u043d\u043d\u0435\u0433\u043e \u043e\u0441\u0442\u0430\u043d\u043e\u0432\u0430 5000 \u0438\u0441\u0445\u043e\u0434\u043d\u044b\u0445 train-\u0437\u0430\u043f\u0438\u0441\u0435\u0439 \u0434\u0435\u0442\u0435\u0440\u043c\u0438\u043d\u0438\u0440\u043e\u0432\u0430\u043d\u043d\u043e \u0434\u0435\u043b\u044f\u0442\u0441\u044f \u043d\u0430 "
        f"{results.get('train_rows', 4000)} \u043e\u0431\u0443\u0447\u0430\u044e\u0449\u0438\u0445 \u0438 {results.get('val_rows', 1000)} validation-\u0437\u0430\u043f\u0438\u0441\u0435\u0439. "
        f"\u041f\u043e\u0441\u043b\u0435 \u043a\u0430\u0436\u0434\u043e\u0439 \u044d\u043f\u043e\u0445\u0438 \u0432\u044b\u0447\u0438\u0441\u043b\u044f\u0435\u0442\u0441\u044f validation-\u043c\u0435\u0442\u0440\u0438\u043a\u0430; \u0435\u0441\u043b\u0438 \u043e\u043d\u0430 \u043d\u0435 \u0443\u043b\u0443\u0447\u0448\u0430\u0435\u0442\u0441\u044f \u043d\u0430 {results.get('min_delta', 0.0001)} "
        f"\u0432 \u0442\u0435\u0447\u0435\u043d\u0438\u0435 {results.get('patience', 2)} \u044d\u043f\u043e\u0445, \u043e\u0431\u0443\u0447\u0435\u043d\u0438\u0435 \u043e\u0441\u0442\u0430\u043d\u0430\u0432\u043b\u0438\u0432\u0430\u0435\u0442\u0441\u044f, \u0430 \u0432\u0435\u0441\u0430 \u043e\u0442\u043a\u0430\u0442\u044b\u0432\u0430\u044e\u0442\u0441\u044f \u043a \u043b\u0443\u0447\u0448\u0435\u0439 \u044d\u043f\u043e\u0445\u0435. "
        "\u0054est-\u0432\u044b\u0431\u043e\u0440\u043a\u0430 \u043f\u0440\u0438\u043c\u0435\u043d\u044f\u0435\u0442\u0441\u044f \u0442\u043e\u043b\u044c\u043a\u043e \u0434\u043b\u044f \u0444\u0438\u043d\u0430\u043b\u044c\u043d\u043e\u0439 \u043e\u0446\u0435\u043d\u043a\u0438.",
    )
    config_rows = pd.DataFrame(
        [
            ["Базовая модель", results["model"]["name"]],
            ["Архитектура", results["model"]["architecture"]],
            ["Режим обучения", results["model"]["training_mode"]],
            ["Projection dim", config["projection_dim"]],
            ["Temperature", config["temperature"]],
            ["Contrastive margin", config["contrastive_margin"]],
            ["Triplet margin", config["triplet_margin"]],
            ["Circle margin / gamma", f"{config['circle_margin']} / {config['circle_gamma']}"],
        ],
        columns=["Параметр", "Значение"],
    )
    add_table(document, "Таблица 2 - Конфигурация эксперимента", config_rows, ["Параметр", "Значение"])


def add_hyperparameter_section(document: Document, search: pd.DataFrame) -> None:
    heading(document, "Предварительный подбор гиперпараметров")
    paragraph(
        document,
        "Перед финальным обучением для каждого сочетания датасета и функции потерь выполнялся отдельный подбор "
        "гиперпараметров. Последовательно варьировались epochs, learning_rate, weight_decay и один параметр, "
        "специфичный для функции потерь. В каждой таблице приведены три проверенных значения; строка со звездочкой "
        "соответствует лучшему значению по основной test-метрике данного сочетания.",
    )

    ordered = search.sort_values(["dataset_id", "loss", "hyperparameter", "value"])
    for (dataset_id, loss_name), combo_frame in ordered.groupby(["dataset_id", "loss"], sort=True):
        combo_title(document, f"{short_dataset(dataset_id)} / {loss_name}")
        for hyperparameter, frame in combo_frame.groupby("hyperparameter", sort=True):
            table_frame = frame.copy()
            best_index = table_frame["primary_metric"].idxmax()
            table_frame["best"] = ["*" if index == best_index else "" for index in table_frame.index]
            table_frame["value"] = table_frame["value"].map(format_value)
            table_frame["metric"] = table_frame["primary_metric"].map(lambda value: f"{value:.3f}")
            table_frame["test loss"] = table_frame["final_test_loss"].map(lambda value: f"{value:.3f}")
            table_frame["train loss"] = table_frame["final_train_loss"].map(lambda value: f"{value:.3f}")
            add_table(
                document,
                f"Подбор {hyperparameter}: {short_dataset(dataset_id)} / {loss_name}",
                table_frame,
                ["best", "value", "metric", "test loss", "train loss"],
            )


def add_results_section(document: Document, results: dict[str, Any], metrics: pd.DataFrame) -> None:
    heading(document, "Результаты")
    paragraph(
        document,
        "Основная метрика зависит от формата обучения: для contrastive pair используется ROC-AUC, для triplet - доля "
        "триплетов с большей близостью positive, для InfoNCE/NT-Xent - top-1 retrieval, для SupCon/Circle - accuracy "
        "1-NN классификации по эмбеддингам.",
    )
    table_metrics = metrics.copy()
    table_metrics["dataset"] = table_metrics["dataset_id"].map(short_dataset)
    table_metrics["metric"] = table_metrics["primary_metric_name"].map(short_metric)
    table_metrics["value"] = table_metrics["primary_metric"].map(lambda value: f"{value:.3f}")
    table_metrics["val metric"] = table_metrics["best_val_metric"].map(lambda value: f"{value:.3f}")
    table_metrics["train metric"] = table_metrics["final_train_metric"].map(lambda value: f"{value:.3f}")
    table_metrics["best epoch"] = table_metrics["best_epoch"].astype(str)
    table_metrics["val loss"] = table_metrics["final_val_loss"].map(lambda value: f"{value:.3f}")
    table_metrics["train loss"] = table_metrics["final_train_loss"].map(lambda value: f"{value:.3f}")
    add_table(
        document,
        "Таблица 3 - Итоговые метрики по функциям потерь",
        table_metrics,
        ["dataset", "loss", "metric", "value", "best epoch", "val metric", "train metric", "val loss", "train loss"],
    )

    baseline_rows = build_baseline_rows(results)
    add_table(
        document,
        "Таблица 4 - Baseline без обучения projection head",
        baseline_rows,
        ["dataset", "view", "metric", "value"],
        max_rows=18,
    )

    add_picture(document, ARTIFACTS_DIR / "plots" / "primary_metrics.png", "Рисунок 1 - Основные метрики по датасетам и функциям потерь")
    add_picture(
        document,
        ARTIFACTS_DIR / "plots" / "loss_panels_by_loss_contrastive.png",
        "\u0420\u0438\u0441\u0443\u043d\u043e\u043a 2 - Contrastive Loss: train/val loss \u0438 \u0446\u0435\u043b\u0435\u0432\u0430\u044f \u043c\u0435\u0442\u0440\u0438\u043a\u0430 \u043f\u043e \u0434\u0430\u0442\u0430\u0441\u0435\u0442\u0430\u043c",
    )
    add_picture(
        document,
        ARTIFACTS_DIR / "plots" / "loss_panels_by_loss_triplet.png",
        "\u0420\u0438\u0441\u0443\u043d\u043e\u043a 3 - Triplet Loss: train/val loss \u0438 \u0446\u0435\u043b\u0435\u0432\u0430\u044f \u043c\u0435\u0442\u0440\u0438\u043a\u0430 \u043f\u043e \u0434\u0430\u0442\u0430\u0441\u0435\u0442\u0430\u043c",
    )
    add_picture(
        document,
        ARTIFACTS_DIR / "plots" / "loss_panels_by_loss_infonce.png",
        "\u0420\u0438\u0441\u0443\u043d\u043e\u043a 4 - InfoNCE Loss: train/val loss \u0438 \u0446\u0435\u043b\u0435\u0432\u0430\u044f \u043c\u0435\u0442\u0440\u0438\u043a\u0430 \u043f\u043e \u0434\u0430\u0442\u0430\u0441\u0435\u0442\u0430\u043c",
    )
    add_picture(
        document,
        ARTIFACTS_DIR / "plots" / "loss_panels_by_loss_nt_xent.png",
        "\u0420\u0438\u0441\u0443\u043d\u043e\u043a 5 - NT-Xent Loss: train/val loss \u0438 \u0446\u0435\u043b\u0435\u0432\u0430\u044f \u043c\u0435\u0442\u0440\u0438\u043a\u0430 \u043f\u043e \u0434\u0430\u0442\u0430\u0441\u0435\u0442\u0430\u043c",
    )
    document.add_page_break()
    add_picture(
        document,
        ARTIFACTS_DIR / "plots" / "loss_panels_by_loss_supcon.png",
        "\u0420\u0438\u0441\u0443\u043d\u043e\u043a 6 - SupCon Loss: train/val loss \u0438 \u0446\u0435\u043b\u0435\u0432\u0430\u044f \u043c\u0435\u0442\u0440\u0438\u043a\u0430 \u043f\u043e \u0434\u0430\u0442\u0430\u0441\u0435\u0442\u0430\u043c",
    )
    add_picture(
        document,
        ARTIFACTS_DIR / "plots" / "loss_panels_by_loss_circle.png",
        "\u0420\u0438\u0441\u0443\u043d\u043e\u043a 7 - Circle Loss: train/val loss \u0438 \u0446\u0435\u043b\u0435\u0432\u0430\u044f \u043c\u0435\u0442\u0440\u0438\u043a\u0430 \u043f\u043e \u0434\u0430\u0442\u0430\u0441\u0435\u0442\u0430\u043c",
    )

    paragraph(
        document,
        "Для all-nli и quora-duplicates наиболее естественными являются pair, triplet и in-batch retrieval постановки. "
        "Для banking77 и twenty-newsgroups class-label losses показывают смысловую структуру классов через 1-NN accuracy. "
        "Низкие значения retrieval_top1 у InfoNCE/NT-Xent на class-derived positive_pair не означают провал обучения: "
        "в этих датасетах у одного класса много равноправных positive-примеров, а метрика требует попасть ровно в пару "
        "с тем же индексом.",
    )


def add_conclusion(document: Document, metrics: pd.DataFrame) -> None:
    _ = metrics
    heading(document, "\u0412\u044b\u0432\u043e\u0434\u044b")
    paragraph(
        document,
        "\u0418\u0442\u043e\u0433\u043e\u0432\u044b\u0435 \u0433\u0440\u0430\u0444\u0438\u043a\u0438 \u0438 \u0442\u0430\u0431\u043b\u0438\u0446\u0430 \u043c\u0435\u0442\u0440\u0438\u043a \u043f\u043e\u043a\u0430\u0437\u044b\u0432\u0430\u044e\u0442, \u0447\u0442\u043e loss-\u0444\u0443\u043d\u043a\u0446\u0438\u0438 \u043d\u0435\u043b\u044c\u0437\u044f "
        "\u043a\u043e\u0440\u0440\u0435\u043a\u0442\u043d\u043e \u0441\u0440\u0430\u0432\u043d\u0438\u0432\u0430\u0442\u044c \u043f\u043e \u043e\u0434\u043d\u043e\u0439 \u0447\u0438\u0441\u043b\u043e\u0432\u043e\u0439 \u0448\u043a\u0430\u043b\u0435: \u0434\u043b\u044f \u0440\u0430\u0437\u043d\u044b\u0445 \u0444\u043e\u0440\u043c\u0430\u0442\u043e\u0432 \u0434\u0430\u043d\u043d\u044b\u0445 "
        "\u043e\u0441\u043d\u043e\u0432\u043d\u044b\u043c\u0438 \u043c\u0435\u0442\u0440\u0438\u043a\u0430\u043c\u0438 \u0431\u044b\u043b\u0438 pair AUC, triplet accuracy, retrieval top-1 \u0438 1-NN accuracy. "
        "\u041f\u043e\u044d\u0442\u043e\u043c\u0443 \u0433\u043b\u0430\u0432\u043d\u044b\u0439 \u0432\u044b\u0432\u043e\u0434 \u043d\u0435 \u0432 \u0442\u043e\u043c, \u0447\u0442\u043e \u043e\u0434\u0438\u043d loss \u00ab\u043b\u0443\u0447\u0448\u0435\u00bb \u0432\u0441\u0435\u0445, \u0430 \u0432 \u0442\u043e\u043c, "
        "\u043a\u0430\u043a\u0430\u044f \u0444\u0443\u043d\u043a\u0446\u0438\u044f \u043f\u043e\u0442\u0435\u0440\u044c \u0441\u043e\u043e\u0442\u0432\u0435\u0442\u0441\u0442\u0432\u0443\u0435\u0442 \u0437\u0430\u0434\u0430\u0447\u0435 \u0438 \u043e\u0440\u0433\u0430\u043d\u0438\u0437\u0430\u0446\u0438\u0438 labels."
    )
    paragraph(
        document,
        "Contrastive Loss \u043b\u0443\u0447\u0448\u0435 \u0432\u0441\u0435\u0433\u043e \u0438\u043d\u0442\u0435\u0440\u043f\u0440\u0435\u0442\u0438\u0440\u0443\u0435\u0442\u0441\u044f \u043a\u0430\u043a loss \u0434\u043b\u044f binary pair verification: "
        "\u043f\u043e\u0438\u0441\u043a \u0434\u0443\u0431\u043b\u0438\u043a\u0430\u0442\u043e\u0432, same/different \u043f\u0430\u0440\u044b, entailment/contradiction \u0438 \u0434\u0440\u0443\u0433\u0438\u0435 \u0441\u0446\u0435\u043d\u0430\u0440\u0438\u0438, "
        "\u0433\u0434\u0435 \u043d\u0443\u0436\u0435\u043d \u043f\u043e\u0440\u043e\u0433 \u0441\u0445\u043e\u0436\u0435\u0441\u0442\u0438. \u0412 \u044d\u0442\u043e\u043c \u0444\u043e\u0440\u043c\u0430\u0442\u0435 \u043e\u043d \u0434\u0430\u0435\u0442 \u0441\u0438\u043b\u044c\u043d\u044b\u0435 test pair AUC: "
        "0.757 \u043d\u0430 NLI, 0.877 \u043d\u0430 Quora, 0.972 \u043d\u0430 Banking77 \u0438 0.826 \u043d\u0430 20News. "
        "\u041d\u0430 \u0433\u0440\u0430\u0444\u0438\u043a\u0430\u0445 \u0443 \u043d\u0435\u0433\u043e \u0447\u0430\u0441\u0442\u043e \u0432\u0438\u0434\u0435\u043d \u0440\u0430\u0437\u0440\u044b\u0432 \u043c\u0435\u0436\u0434\u0443 train \u0438 val \u043c\u0435\u0442\u0440\u0438\u043a\u043e\u0439, \u043f\u043e\u044d\u0442\u043e\u043c\u0443 "
        "\u0435\u0433\u043e \u0441\u0442\u043e\u0438\u0442 \u0438\u0441\u043f\u043e\u043b\u044c\u0437\u043e\u0432\u0430\u0442\u044c \u0441 early stopping \u043f\u043e validation-\u043c\u0435\u0442\u0440\u0438\u043a\u0435."
    )
    paragraph(
        document,
        "Triplet Loss \u043d\u0430\u0438\u0431\u043e\u043b\u0435\u0435 \u043d\u0430\u0433\u043b\u044f\u0434\u0435\u043d \u0434\u043b\u044f \u0437\u0430\u0434\u0430\u0447 \u0440\u0430\u043d\u0436\u0438\u0440\u043e\u0432\u0430\u043d\u0438\u044f \u0438 \u043e\u0442\u043d\u043e\u0441\u0438\u0442\u0435\u043b\u044c\u043d\u043e\u0439 "
        "\u0431\u043b\u0438\u0437\u043e\u0441\u0442\u0438, \u0433\u0434\u0435 \u0435\u0441\u0442\u044c \u043e\u0441\u043c\u044b\u0441\u043b\u0435\u043d\u043d\u0430\u044f \u0442\u0440\u043e\u0439\u043a\u0430 anchor-positive-negative. \u0418\u0442\u043e\u0433\u043e\u0432\u044b\u0435 \u0433\u0440\u0430\u0444\u0438\u043a\u0438 "
        "\u043f\u043e\u043a\u0430\u0437\u044b\u0432\u0430\u044e\u0442, \u0447\u0442\u043e best epoch \u0434\u043b\u044f \u043d\u0435\u0433\u043e \u043e\u0431\u044b\u0447\u043d\u043e \u043d\u0430\u0441\u0442\u0443\u043f\u0430\u0435\u0442 \u0440\u0430\u043d\u043e, \u043d\u043e test triplet accuracy "
        "\u043e\u0441\u0442\u0430\u0435\u0442\u0441\u044f \u0432\u044b\u0441\u043e\u043a\u043e\u0439: 0.866 \u043d\u0430 NLI, 0.954 \u043d\u0430 Quora, 0.974 \u043d\u0430 Banking77 \u0438 0.827 \u043d\u0430 20News. "
        "\u042d\u0442\u043e \u0434\u0435\u043b\u0430\u0435\u0442 Triplet Loss \u0441\u0438\u043b\u044c\u043d\u044b\u043c \u0432\u044b\u0431\u043e\u0440\u043e\u043c \u0434\u043b\u044f retrieval/ranking-\u0441\u0446\u0435\u043d\u0430\u0440\u0438\u0435\u0432, \u043d\u043e \u0442\u0440\u0435\u0431\u0443\u0435\u0442 "
        "\u043a\u0430\u0447\u0435\u0441\u0442\u0432\u0435\u043d\u043d\u043e\u0433\u043e \u0441\u043f\u043e\u0441\u043e\u0431\u0430 \u0441\u0435\u043c\u043f\u043b\u0438\u0440\u043e\u0432\u0430\u0442\u044c negatives."
    )
    paragraph(
        document,
        "InfoNCE \u0438 NT-Xent \u043f\u043e\u0434\u0445\u043e\u0434\u044f\u0442 \u0434\u043b\u044f in-batch retrieval, \u0433\u0434\u0435 \u0443 anchor \u0435\u0441\u0442\u044c \u043e\u0434\u0438\u043d \u0446\u0435\u043b\u0435\u0432\u043e\u0439 "
        "positive \u0432 \u0431\u0430\u0442\u0447\u0435. \u041d\u0430 NLI \u0438 Quora \u044d\u0442\u0438 loss \u0434\u0430\u044e\u0442 \u0430\u0434\u0435\u043a\u0432\u0430\u0442\u043d\u044b\u0439 test retrieval top-1: "
        "\u0434\u043b\u044f InfoNCE 0.769 \u0438 0.853, \u0434\u043b\u044f NT-Xent 0.758 \u0438 0.848. \u041d\u0430 Banking77 \u0438 20News \u043e\u043d\u0438 "
        "\u0432\u044b\u0433\u043b\u044f\u0434\u044f\u0442 \u043f\u043b\u043e\u0445\u043e \u0438\u043c\u0435\u043d\u043d\u043e \u0432 exact paired-index \u043c\u0435\u0442\u0440\u0438\u043a\u0435: test top-1 \u0441\u043d\u0438\u0436\u0430\u0435\u0442\u0441\u044f \u0434\u043e "
        "0.019/0.005 \u0443 InfoNCE \u0438 0.018/0.002 \u0443 NT-Xent. \u042d\u0442\u043e \u043d\u0435 \u0434\u043e\u043a\u0430\u0437\u044b\u0432\u0430\u0435\u0442 \u043d\u0435\u043f\u0440\u0438\u0433\u043e\u0434\u043d\u043e\u0441\u0442\u044c "
        "\u0441\u0430\u043c\u0438\u0445 loss-\u0444\u0443\u043d\u043a\u0446\u0438\u0439, \u0430 \u043f\u043e\u043a\u0430\u0437\u044b\u0432\u0430\u0435\u0442 \u043d\u0435\u0441\u043e\u0432\u043f\u0430\u0434\u0435\u043d\u0438\u0435 loss, \u0440\u0430\u0437\u043c\u0435\u0442\u043a\u0438 \u0438 \u043c\u0435\u0442\u0440\u0438\u043a\u0438: "
        "\u043d\u0430 class-label \u0434\u0430\u043d\u043d\u044b\u0445 \u0443 \u043e\u0434\u043d\u043e\u0433\u043e anchor \u0435\u0441\u0442\u044c \u043c\u043d\u043e\u0433\u043e \u0440\u0430\u0432\u043d\u043e\u043f\u0440\u0430\u0432\u043d\u044b\u0445 positives."
    )
    paragraph(
        document,
        "SupCon \u0438 Circle Loss \u0441\u0442\u043e\u0438\u0442 \u0431\u0440\u0430\u0442\u044c \u0434\u043b\u044f class-label \u0434\u0430\u043d\u043d\u044b\u0445: intent classification, topic clustering "
        "\u0438 \u0434\u0440\u0443\u0433\u0438\u0445 \u0441\u0446\u0435\u043d\u0430\u0440\u0438\u0435\u0432, \u0433\u0434\u0435 \u043d\u0443\u0436\u043d\u043e \u0441\u0431\u043b\u0438\u0437\u0438\u0442\u044c \u0432\u0441\u0435 \u043e\u0431\u044a\u0435\u043a\u0442\u044b \u043e\u0434\u043d\u043e\u0433\u043e \u043a\u043b\u0430\u0441\u0441\u0430. "
        "\u041d\u0430 Banking77 Circle Loss \u0434\u0430\u0435\u0442 \u043b\u0443\u0447\u0448\u0438\u0439 class-label test \u0440\u0435\u0437\u0443\u043b\u044c\u0442\u0430\u0442: 1-NN accuracy 0.903 "
        "\u043f\u0440\u043e\u0442\u0438\u0432 0.879 \u0443 SupCon. \u041d\u0430 20News test-\u0440\u0430\u0437\u043d\u0438\u0446\u0430 \u043c\u0430\u043b\u0430: 0.611 \u0443 SupCon \u0438 0.609 \u0443 Circle, "
        "\u043d\u043e \u043d\u0430 validation-\u0433\u0440\u0430\u0444\u0438\u043a\u0435 Circle \u0432\u044b\u0448\u0435 (0.672 \u043f\u0440\u043e\u0442\u0438\u0432 0.648). \u0414\u043b\u044f \u0441\u0438\u043b\u044c\u043d\u043e \u0440\u0430\u0437\u043c\u0435\u0447\u0435\u043d\u043d\u044b\u0445 "
        "\u043a\u043b\u0430\u0441\u0441\u043e\u0432 \u043f\u043e \u044d\u0442\u0438\u043c \u0433\u0440\u0430\u0444\u0438\u043a\u0430\u043c Circle \u0432\u044b\u0433\u043b\u044f\u0434\u0438\u0442 \u0447\u0443\u0442\u044c \u0431\u043e\u043b\u0435\u0435 \u0443\u0441\u0442\u043e\u0439\u0447\u0438\u0432\u044b\u043c, \u0430 SupCon - "
        "\u0431\u043e\u043b\u0435\u0435 \u043f\u0440\u043e\u0441\u0442\u044b\u043c \u0438 \u0431\u043b\u0438\u0437\u043a\u0438\u043c \u043f\u043e \u043a\u0430\u0447\u0435\u0441\u0442\u0432\u0443 \u0432\u0430\u0440\u0438\u0430\u043d\u0442\u043e\u043c."
    )
    paragraph(
        document,
        "\u041f\u0440\u0430\u043a\u0442\u0438\u0447\u0435\u0441\u043a\u0430\u044f \u043a\u0430\u0440\u0442\u0430 \u0432\u044b\u0431\u043e\u0440\u0430 \u043f\u043e \u0438\u0442\u043e\u0433\u0430\u043c \u044d\u043a\u0441\u043f\u0435\u0440\u0438\u043c\u0435\u043d\u0442\u0430 \u0442\u0430\u043a\u0430\u044f: \u0434\u043b\u044f pair verification - "
        "Contrastive Loss; \u0434\u043b\u044f ranking/retrieval \u0441 \u044f\u0432\u043d\u044b\u043c negative - Triplet Loss; \u0434\u043b\u044f in-batch retrieval "
        "\u0441 \u043e\u0434\u043d\u0438\u043c \u0446\u0435\u043b\u0435\u0432\u044b\u043c positive - InfoNCE \u0438 NT-Xent; \u0434\u043b\u044f \u0434\u0430\u043d\u043d\u044b\u0445 \u0441 \u043c\u0435\u0442\u043a\u0430\u043c\u0438 \u043a\u043b\u0430\u0441\u0441\u043e\u0432 - "
        "SupCon \u0438 Circle Loss. \u0412\u043e \u0432\u0441\u0435\u0445 \u0441\u043b\u0443\u0447\u0430\u044f\u0445 best epoch \u043d\u0443\u0436\u043d\u043e \u0432\u044b\u0431\u0438\u0440\u0430\u0442\u044c \u043f\u043e \u0446\u0435\u043b\u0435\u0432\u043e\u0439 validation-\u043c\u0435\u0442\u0440\u0438\u043a\u0435: "
        "\u0438\u0442\u043e\u0433\u043e\u0432\u044b\u0435 \u0433\u0440\u0430\u0444\u0438\u043a\u0438 \u043f\u043e\u043a\u0430\u0437\u044b\u0432\u0430\u044e\u0442, \u0447\u0442\u043e train loss \u043c\u043e\u0436\u0435\u0442 \u043f\u0430\u0434\u0430\u0442\u044c \u0434\u0430\u043b\u044c\u0448\u0435 \u0431\u0435\u0437 \u0443\u043b\u0443\u0447\u0448\u0435\u043d\u0438\u044f "
        "\u0446\u0435\u043b\u0435\u0432\u043e\u0439 \u043c\u0435\u0442\u0440\u0438\u043a\u0438 \u043d\u0430 validation."
    )

def add_sources(document: Document) -> None:
    heading(document, "Список использованных источников")
    sources = [
        "Hadsell R., Chopra S., LeCun Y. Dimensionality Reduction by Learning an Invariant Mapping. 2006.",
        "Hoffer E., Ailon N. Deep Metric Learning using Triplet Network. 2014.",
        "van den Oord A., Li Y., Vinyals O. Representation Learning with Contrastive Predictive Coding. 2018.",
        "Chen T., Kornblith S., Norouzi M., Hinton G. A Simple Framework for Contrastive Learning of Visual Representations. 2020.",
        "Khosla P. et al. Supervised Contrastive Learning. 2020.",
        "Sun Y. et al. Circle Loss: A Unified Perspective of Pair Similarity Optimization. 2020.",
        "Reimers N., Gurevych I. Sentence-BERT: Sentence Embeddings using Siamese BERT-Networks. 2019.",
        "Hugging Face model card: sentence-transformers/msmarco-distilbert-base-dot-prod-v3.",
        "Hugging Face datasets: sentence-transformers/all-nli, sentence-transformers/quora-duplicates, mteb/banking77, SetFit/20_newsgroups, sentence-transformers/stsb.",
    ]
    for index, source in enumerate(sources, start=1):
        paragraph(document, f"{index}. {source}")


def build_baseline_rows(results: dict[str, Any]) -> pd.DataFrame:
    rows = []
    for dataset_id, views in results["baseline"].items():
        if dataset_id == "stsb_eval":
            rows.append([short_dataset(dataset_id), "STS", "Spearman", round(views["spearman"], 3)])
            continue
        for view, metrics in views.items():
            if "pair_auc" in metrics:
                rows.append([short_dataset(dataset_id), short_view(view), "Pair AUC", round(metrics["pair_auc"], 3)])
            elif "triplet_accuracy" in metrics:
                rows.append([short_dataset(dataset_id), short_view(view), "Triplet acc.", round(metrics["triplet_accuracy"], 3)])
            elif "retrieval_top1" in metrics:
                rows.append([short_dataset(dataset_id), short_view(view), "Top-1", round(metrics["retrieval_top1"], 3)])
            elif "knn1_accuracy" in metrics:
                rows.append([short_dataset(dataset_id), short_view(view), "1-NN acc.", round(metrics["knn1_accuracy"], 3)])
    return pd.DataFrame(rows, columns=["dataset", "view", "metric", "value"])


def add_table(
    document: Document,
    caption: str,
    frame: pd.DataFrame,
    columns: list[str],
    max_rows: int | None = None,
    font_size: int = 12,
) -> None:
    paragraph(document, caption)
    table_frame = frame[columns].head(max_rows) if max_rows else frame[columns]
    table = document.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, column in zip(table.rows[0].cells, columns, strict=True):
        cell.text = str(column)
    repeat_table_header(table.rows[0])
    for row in table_frame.itertuples(index=False):
        cells = table.add_row().cells
        for cell, value in zip(cells, row, strict=True):
            cell.text = str(value)
    set_table_font_size(table, font_size)
    add_spacer(document, 1)


def repeat_table_header(row: Any) -> None:
    table_row_properties = row._tr.get_or_add_trPr()
    if table_row_properties.find(qn("w:tblHeader")) is None:
        table_row_properties.append(OxmlElement("w:tblHeader"))


def set_table_font_size(table: Any, font_size: int) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph_item in cell.paragraphs:
                for run in paragraph_item.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(font_size)


def short_dataset(dataset_id: str) -> str:
    return {
        "all_nli": "NLI",
        "quora_duplicates": "Quora",
        "banking77": "Banking77",
        "twenty_newsgroups": "20News",
        "stsb_eval": "STS-B",
    }.get(dataset_id, dataset_id)


def short_metric(metric: str) -> str:
    return {
        "pair_auc": "Pair AUC",
        "triplet_accuracy": "Triplet acc.",
        "retrieval_top1": "Top-1",
        "knn1_accuracy": "1-NN acc.",
    }.get(metric, metric)


def short_view(view: str) -> str:
    return {
        "contrastive_pair": "Pair",
        "triplet": "Triplet",
        "positive_pair": "Pos. pair",
        "class_text": "Class",
        "sts_pair": "STS",
    }.get(view, view)


def format_value(value: Any) -> str:
    try:
        numeric = float(value)
    except (TypeError, ValueError):
        return str(value)
    if numeric.is_integer():
        return str(int(numeric))
    return f"{numeric:g}"


def combo_title(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)


def add_picture(document: Document, path: Path, caption: str) -> None:
    document.add_picture(str(path), width=Inches(6.5))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    centered(document, caption)


def heading(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 0, 0)


def paragraph(document: Document, text: str) -> None:
    p = document.add_paragraph(text)
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def centered(document: Document, text: str, bold: bool = False) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)


def small_centered(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)


def add_spacer(document: Document, count: int) -> None:
    for _ in range(count):
        document.add_paragraph("")


if __name__ == "__main__":
    build_report()
